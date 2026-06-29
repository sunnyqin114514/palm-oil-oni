# -*- coding: utf-8 -*-
"""生成"温度深化研究 + 三维产量预测模型 + 2024-2026 样本外验证"Word 报告。

覆盖用户三条研究方向:
  1. 温度对产量关系: 积温窗口长度优化 / 高温低温非对称 / 温度×降水空间交互
  2. ONI+降水+温度 三维产量预测模型
  3. 模型参数基于 2010-2023, 用 2024-2026 预测与实际对比验证

只读取既有分析产物 (temp_research/*.csv, model3d_*.csv/json, 各图), 不重跑模型。

输出: reports/Temperature_3D_Model_Report.docx
"""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
REPORTS_DIR = PROJECT_ROOT / "reports"
MODEL_DIR = PROJECT_ROOT / "code" / "model"
TEMP_DIR = MODEL_DIR / "temp_research"
FIG_DIR = MODEL_DIR / "figures"
METEO_DIR = PROJECT_ROOT / "data" / "processed" / "meteo"
OUT_DOCX = REPORTS_DIR / "Temperature_3D_Model_Report.docx"


def add_para(doc, text="", *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic, r.font.size = bold, italic, Pt(size)
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10.5)
    return p


def add_formula(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    return p


def make_table(doc, headers, rows, font_size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        run = c.paragraphs[0].add_run(str(h))
        run.bold = True
        run.font.size = Pt(font_size)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            cells[i].paragraphs[0].add_run(str(v)).font.size = Pt(font_size)
    return t


def add_pic(doc, path: Path, caption: str, width=6.0):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = add_para(doc, caption, italic=True, size=9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[图缺失] {path.name}", italic=True, size=9)


def f(x, d=4):
    try:
        if pd.isna(x):
            return "-"
        return f"{float(x):.{d}f}"
    except (TypeError, ValueError):
        return str(x)


def main() -> None:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    dim_a = pd.read_csv(TEMP_DIR / "dim_a_window.csv")
    dim_b = pd.read_csv(TEMP_DIR / "dim_b_asymmetry.csv")
    dim_c = pd.read_csv(TEMP_DIR / "dim_c_spatial.csv")
    metrics = pd.read_csv(MODEL_DIR / "model3d_validation_metrics.csv")
    weights = json.loads((MODEL_DIR / "model3d_weights.json").read_text(encoding="utf-8"))
    reg = pd.read_csv(METEO_DIR / "malaysia_regional_weather_monthly.csv")

    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(11)

    # ---------- 封面 ----------
    title = doc.add_heading("棕榈油单产研究：温度深化分析与 ONI-降水-温度三维预测模型", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = add_para(doc, "含 2024-2026 样本外验证 · 全马来西亚格点气象 (含婆罗洲)", italic=True, size=11)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "")

    # ---------- 摘要 ----------
    add_heading(doc, "一、研究目标与核心结论", level=1)
    add_para(doc, "本轮针对三条研究方向展开：", size=11)
    add_bullet(doc, "①温度对产量的关系（优化积温窗口长度、区分高温/低温、温度×降水空间交互）；")
    add_bullet(doc, "②建立 ONI+降水+温度三维产量预测模型；")
    add_bullet(doc, "③模型参数基于 2010-2023 估计，用 2024-2026 的预测与真实值对比验证。")
    add_para(doc, "核心结论：", bold=True)
    add_bullet(doc, "温度的『主效应』较弱：单看全国温度距平，各积温窗口在样本内均不显著（p≈0.10~0.57），"
                    "10 个月窗口最强但仅 p=0.10。")
    add_bullet(doc, "温度效应『怕热不怕冷』且呈非对称：偏热部分系数为负、接近显著（p≈0.10），"
                    "偏冷部分不显著（p≈0.35）——高于常态的温度才压制单产。")
    add_bullet(doc, "温度真正的作用通过『空间交互』显形：西马（半岛）温度×降水交互项显著（p=0.023），"
                    "捕捉『又热又旱』复合胁迫；砂拉越、沙巴的交互项不显著。")
    bm, bv = weights["model_name"], weights["variant"]
    mt = weights["metrics"]
    add_bullet(doc, f"三维模型 M2（ONI+降水+温度+西马交互）样本外表现最佳："
                    f"2024-2026 RMSE={f(mt['test_rmse'],4)}、MAPE={f(mt['test_mape_pct'],2)}%、"
                    f"方向命中={f(mt['test_dir_hit']*100,1)}%、样本外 R²={f(mt['test_r2'],3)}。")

    # ---------- 数据基础 ----------
    add_heading(doc, "二、数据基础的扩展", level=1)
    add_para(doc, "为支撑空间交互与样本外验证，本轮补齐三类数据：")
    add_bullet(doc, "全马来西亚格点气象：原有 NASA POWER 格点经度只到 109.375°E（仅覆盖半岛+南海），"
                    "婆罗洲（沙巴 115-119°E、砂拉越 109.6-115°E）几乎缺失。本轮重新抓取覆盖 "
                    "lon 99.375-119.5°E / lat 0.5-7.5°N 的全马格点（495 格点，2007-01~2026-05）。")
    add_bullet(doc, "2026 年气象：NASA POWER 月度产品仅到 2025-12，2026 改用日度数据聚合到月"
                    "（气温取月内日均、降水取月内日总），覆盖到 2026-05。口径已与历史国家序列校验一致"
                    "（如 PRCP 2025-01=287.074、T2M 2025-01=26.229，分毫不差）。")
    add_bullet(doc, "三区气象序列：用陆地掩膜（global_land_mask）剔除海面格点后，按经纬度框划分西马/砂拉越/沙巴三区，"
                    "各区取陆地格点平均。")
    reg_means = reg.groupby("Region").agg(T2M=("T2M", "mean"), PRCP=("PRCP", "mean"),
                                          n=("n_cells", "first")).reset_index()
    make_table(doc, ["区域", "陆地格点数", "气温年均(°C)", "降水年均(mm/月)"],
               [[r.Region, int(r.n), f(r.T2M, 2), f(r.PRCP, 1)] for r in reg_means.itertuples()])
    add_para(doc, "可见三区气候确有差异：砂拉越最湿（约 306 mm/月），婆罗洲两区较半岛偏凉约 1.2°C，"
                  "为空间交互分析提供了真实的空间异质性。", size=10)
    add_para(doc, "面积口径：2015-2025 用真实 MPOB 成熟面积；2026 用近三年线性趋势外推（约 4.99M ha，标记 FORECAST）；"
                  "2010-2014 用全样本线性回推（标记 BACKCAST，下文以『含/不含回推』两套对比其影响）。", size=10)

    # ---------- 维度 a ----------
    add_heading(doc, "三、温度对产量的关系", level=1)
    add_heading(doc, "3.1 维度 a：积温（累计温度）窗口长度优化", level=2)
    add_para(doc, "做法：固定 const+Trend+月份哑变量与基础项（ONI_lag12 + 3 月滚动降水距平），"
                  "把温度距平的滚动窗口从 1 个月搜索到 12 个月，看哪个窗口的温度信号最显著、样本外预测最好。")
    rows_a = [[int(r.window_months), f(r.adj_r2, 4), f(r.focus_beta, 5), f(r.focus_p, 4),
               f(r.test_rmse, 5)] for r in dim_a.itertuples()]
    make_table(doc, ["温度积温窗口(月)", "训练调整R²", "温度系数β", "温度p值", "样本外RMSE"], rows_a, font_size=8.5)
    add_para(doc, "结论：所有窗口温度系数均为负（温度偏高→单产偏低，符合热胁迫预期）；"
                  "样本内显著性随窗口加长而上升，10 个月窗口最强（p=0.098）；但样本外 RMSE 在短窗口略低，"
                  "说明长窗口温度信号与 12 个月 ONI 滞后存在部分重叠。综合取 10 个月作为温度积温窗口。", size=10)
    add_pic(doc, TEMP_DIR / "dim_a_window.png", "图1 积温窗口长度 vs 样本外RMSE(红) / 训练显著性(蓝)")

    # ---------- 维度 b ----------
    add_heading(doc, "3.2 维度 b：高温 / 低温的非对称效应", level=2)
    add_para(doc, "做法：以 10 个月窗口为基础，把温度距平拆成偏热部分 max(ΔT,0) 与偏冷部分 min(ΔT,0)，"
                  "分别/共同放入模型，看单产到底对哪一侧敏感。")
    dim_b_idx = dim_b.set_index("model")

    def _bget(model, col):
        return dim_b_idx.loc[model][col] if col in dim_b_idx.columns else None

    rows_b = []
    for model in dim_b_idx.index:
        rows_b.append([model, f(dim_b_idx.loc[model]["adj_r2"], 4), f(dim_b_idx.loc[model]["test_rmse"], 5),
                       f(_bget(model, "beta[TAVG_HEAT_10m]"), 5), f(_bget(model, "p[TAVG_HEAT_10m]"), 4),
                       f(_bget(model, "beta[TAVG_COLD_10m]"), 5), f(_bget(model, "p[TAVG_COLD_10m]"), 4)])
    make_table(doc, ["模型", "训练调整R²", "样本外RMSE", "偏热β", "偏热p", "偏冷β", "偏冷p"], rows_b, font_size=8.5)
    add_para(doc, "结论：偏热部分系数为负且接近显著（β≈-0.037，p≈0.10），偏冷部分不显著（p≈0.35）。"
                  "即棕榈油单产对『高于常态的温度』敏感，对『低于常态的温度』基本不敏感——典型的热胁迫非对称效应。", size=10)
    add_pic(doc, TEMP_DIR / "dim_b_asymmetry.png", "图2 偏热 vs 偏冷 对单产的敏感度（系数）", width=4.6)

    # ---------- 维度 c ----------
    add_heading(doc, "3.3 维度 c：温度 × 降水的空间交互", level=2)
    add_para(doc, "做法：在三因子基础上分别加入『全国 / 三区』的温度×降水交互项与『又热又旱』复合胁迫项，"
                  "看空间交互是否比全国单一交互更有解释/预测力。交互项 = 温度距平 × 降水距平（10 月窗口）。")
    rows_c = [[r.model, r.added_term, f(r.adj_r2, 4), f(r.added_beta, 5), f(r.added_p, 4),
               f(r.test_rmse, 5), f(r.test_dir_hit, 3)] for r in dim_c.itertuples()]
    make_table(doc, ["模型(基础三因子+)", "新增项", "训练调整R²", "新增项β", "新增项p", "样本外RMSE", "方向命中"],
               rows_c, font_size=8)
    add_para(doc, "结论：西马（半岛）温度×降水交互项显著（p=0.023）且样本外 RMSE 从 0.0445 降到 0.0310、"
                  "方向命中从 82% 升到 89%，是最强的空间信号；全国交互次之（p=0.052）；砂拉越、沙巴交互项均不显著。"
                  "交互系数为正意味着：温度与降水反向（又热又旱 / 又冷又涝）时单产走低，其中『又热又旱』是主要复合胁迫。"
                  "这与半岛是成熟高产区、对水热复合胁迫更敏感的农学直觉一致。", size=10)
    add_pic(doc, TEMP_DIR / "dim_c_spatial.png", "图3 各空间交互方案的样本外RMSE（越短越好，灰线为基础三因子）")

    # ---------- 三维模型 ----------
    add_heading(doc, "四、ONI-降水-温度 三维预测模型", level=1)
    add_para(doc, "在统一结构 const+Trend+月份哑变量 之上，比较三套气候设定 × 两套训练样本：")
    add_formula(doc, "M0_baseline : ONI_lag12 + PRCP_DEV_3m + TAVG_DEV_3m   (前期结构, 对照)")
    add_formula(doc, "M1_3factor  : ONI_lag12 + PRCP_DEV_3m + TAVG_DEV_10m  (温度窗口优化)")
    add_formula(doc, "M2_spatial  : M1 + INTX_West_10m   (加西马温度×降水空间交互)")
    add_para(doc, "训练样本两套：full=2010-2023（含 2010-2014 面积回推）；real=2015-2023（仅真实 MPOB 面积）。"
                  "样本外验证窗口统一为 2024-01~2026-05。", size=10)
    order = {"M0_baseline": 0, "M1_3factor": 1, "M2_spatial": 2}
    metrics["k"] = metrics["model"].map(order)
    metrics = metrics.sort_values(["variant", "k"])
    rows_m = [[r.model, r.variant, int(r.n_train), f(r.adj_r2, 4), f(r.test_rmse, 5),
               f(r["test_mape%"], 2), f(r.test_dir_hit, 3), f(r.test_r2, 3)]
              for _, r in metrics.iterrows()]
    make_table(doc, ["模型", "训练样本", "样本量", "训练调整R²", "样本外RMSE", "MAPE%", "方向命中", "样本外R²"],
               rows_m, font_size=8.5)
    add_para(doc, "要点：", bold=True, size=10.5)
    add_bullet(doc, "M2（加西马交互）在两套样本中都最优；仅把温度换成 10 月窗口（M1）反而略差于基线，"
                    "再次印证温度的价值在『交互』而非『主效应』。")
    add_bullet(doc, "含回推的 full（2010-2023, 168 个月）样本外反而优于 real（2015-2023, 108 个月）："
                    "更长的训练窗口让趋势与季节系数估计更稳，2010-2014 面积高估带来的偏差主要被截距/趋势吸收，"
                    "未污染驱动样本外波动的气候/季节系数。即『前推扩样本』对样本外验证利大于弊。")

    # ---------- 样本外验证 ----------
    add_heading(doc, "五、2024-2026 样本外验证（推荐模型 M2_spatial / full）", level=1)
    coef = weights["coef"]
    pv = weights["p_values"]
    add_para(doc, f"推荐模型：{bm}（训练 {weights['trained_window']}）。系数与显著性：", size=10.5)
    feat_rows = [["截距", f(weights["intercept"], 5), "-"]]
    for k in ["Trend", "ONI_lag12", "PRCP_DEV_3m", "TAVG_DEV_10m", "INTX_West_10m"]:
        feat_rows.append([k, f(coef.get(k), 6), f(pv.get(k), 4)])
    make_table(doc, ["项", "系数", "p值"], feat_rows, font_size=9)
    add_para(doc, "样本外指标（2024-01~2026-05, 共 29 个月）：", size=10.5, bold=True)
    add_bullet(doc, f"RMSE={f(mt['test_rmse'],4)}（单位 吨/公顷）、MAE={f(mt['test_mae'],4)}、"
                    f"MAPE={f(mt['test_mape_pct'],2)}%；")
    add_bullet(doc, f"方向命中率={f(mt['test_dir_hit']*100,1)}%（环比涨跌方向）、样本外 R²={f(mt['test_r2'],3)}。")
    add_pic(doc, FIG_DIR / "model3d_validation_timeline.png", "图4 全样本：真实单产 vs 训练期拟合 / 样本外预测")
    add_pic(doc, FIG_DIR / "model3d_oos_detail.png", "图5 样本外 2024-2026：逐月对比（左）与 预测-真实散点（右）")
    add_para(doc, "可见模型能稳定复现 2024-2026 单产的季节起伏与多数转折，散点贴近 45° 线，"
                  "平均相对误差约 7.7%。", size=10)

    # ---------- 结论与局限 ----------
    add_heading(doc, "六、结论、模型定位与局限", level=1)
    add_bullet(doc, "温度对单产的影响是『非对称 + 交互依赖』的：单纯的温度高低解释力有限，"
                    "关键是高温叠加少雨（又热又旱）在半岛主产区造成的复合胁迫。")
    add_bullet(doc, "三维模型应以 ONI(滞后12月) + 降水(3月积水) + 温度(10月积温) 为骨架，"
                    "并加入西马温度×降水交互项，样本外精度显著提升（RMSE 降约 30%）。")
    add_bullet(doc, "样本量权衡：在当前数据下，前推到 2010 训练（即便 2010-2014 面积为线性回推）"
                    "比只用 2015 起的真实样本，样本外更稳。建议以 full 为主、real 为稳健性对照。")
    add_para(doc, "局限：", bold=True, size=10.5)
    add_bullet(doc, "2010-2014 成熟面积为线性回推（真实值未公开），会系统性低估该段单产水平，"
                    "故其结论仅用于趋势/系数稳健性，不用于绝对水平判断；2026 面积为外推、且仅到 5 月。")
    add_bullet(doc, "婆罗洲两区的格点框可能含极少量邻近非马来西亚陆地，代表区域气候而非精确行政边界；"
                    "ONI_lag12 在更长窗口下显著性减弱（p=0.33），但符号稳定、属理论驱动项，予以保留。")
    add_bullet(doc, "样本外窗口仅 29 个月，指标存在小样本波动；后续可随 MPOB/气象数据更新滚动复验。")

    # ---------- 文件清单 ----------
    add_heading(doc, "七、相关文件清单", level=1)
    files = [
        ["code/scripts/fetch_nasa_power_full_malaysia.py", "抓取全马格点气象(含婆罗洲)+2026"],
        ["code/model/regional_weather.py", "三区气象序列(陆地掩膜)"],
        ["code/model/build_extended_dataset.py", "2010-2026 扩展建模数据集"],
        ["code/model/feature_builder.py", "积温/非对称/三区交互候选特征"],
        ["code/model/temperature_analysis.py", "温度三维度分析(a/b/c)"],
        ["code/model/build_3factor_model.py", "三维模型训练 + 样本外验证"],
        ["code/model/model3d_weights.json", "推荐模型权重"],
        ["code/model/temp_research/*.csv/png", "温度三维度结果表与图"],
        ["data/processed/meteo/malaysia_regional_weather_*.csv", "三区气象序列"],
        ["data/processed/product/palm_oil_extended_dataset.csv", "扩展数据集"],
    ]
    make_table(doc, ["文件", "说明"], files, font_size=9)

    doc.save(str(OUT_DOCX))
    print(f"[ok] 报告已生成 -> {OUT_DOCX}")


if __name__ == "__main__":
    main()
