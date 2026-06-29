# -*- coding: utf-8 -*-
"""
生成交付物 R3 / #4：券商研报精华萃取（棕榈油 × 厄尔尼诺）
- 输入：4 份券商研报（已人工阅读，下述内容为转述提炼，非原文复制粘贴）
- 输出：reports/R3_券商研报精华萃取.docx
要求对齐：清单表(标题/机构/日期/链接) + 逐份萃取(结论/数据指标/方法/可借鉴点各≥60字) + ≥300字交叉对比
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
OUT_PATH = os.path.join(PROJECT_ROOT, "reports", "R3_券商研报精华萃取.docx")

EAST_ASIA_FONT = "宋体"
HEAD_FONT = "微软雅黑"


def set_cn_font(run, font=EAST_ASIA_FONT, size=None, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, size=10.5, bold=False, font=EAST_ASIA_FONT,
             align=None, space_after=4, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_cn_font(run, font=font, size=size, bold=bold, color=color)
    return p


def add_labeled(doc, label, body):
    """带蓝色小标题（结论/数据指标/方法/可借鉴点）的段落。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    r1 = p.add_run(f"【{label}】")
    set_cn_font(r1, font=HEAD_FONT, size=10.5, bold=True, color=(0x1F, 0x4E, 0x79))
    r2 = p.add_run(body)
    set_cn_font(r2, font=EAST_ASIA_FONT, size=10.5)
    return p


def add_heading(doc, text, size=13, color=(0x1F, 0x4E, 0x79), space_before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_cn_font(run, font=HEAD_FONT, size=size, bold=True, color=color)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color})
    tcPr.append(shd)


def fill_cell(cell, text, bold=False, size=9.5, color=None, font=EAST_ASIA_FONT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    set_cn_font(run, font=font, size=size, bold=bold, color=color)


# ---------------------------------------------------------------------------
# 数据：四份研报（标题/机构/日期/链接/四类萃取）
# ---------------------------------------------------------------------------
NOTE_LINK = "微信群分发 PDF，无公开链接"

REPORTS = [
    {
        "no": "①",
        "title": "棕榈油周度报告：NOAA 周度报告与棕榈油产地降雨跟踪——赤道太平洋海域热量储备已超 2015 年",
        "org": "上海东证期货（东证衍生品研究院）",
        "date": "2026-06-10",
        "tag": "棕榈油专题 / 气象监测",
        "link": NOTE_LINK,
        "conclusion": (
            "走势评级给出“震荡”。官方状态为 El Niño Watch（厄尔尼诺监测预警），处于 ENSO 中性向"
            "厄尔尼诺的转型期；NOAA 5 月将 5–7 月厄尔尼诺被确认概率上调至 80% 以上，年底强/超强厄尔尼诺"
            "概率约 67%。核心判断：赤道太平洋上层海洋热量储备已超过 2015 年超强厄尔尼诺峰值，能量基础充足，"
            "西风异常全面爆发或在二季度末至三季度引发新一轮下沉开尔文波，带动厄尔尼诺加速发展；6 月初产地已现少雨端倪。"),
        "data": (
            "截至 6 月 3 日当周 Nino1+2/3/3.4/4 海区周度 SSTA（Sea Surface Temperature Anomalies，海表温度距平）"
            "为 2.6/1.5/1.3/1.1℃，相对 SSTA（RONI 口径）为 2.1/1/0.7/0.7℃；温跃层斜率降至约 -30 米"
            "（2015/16 强厄尔尼诺约 -60 米）；0–300m 上层海洋热量异常接近 +2℃；2026 年 3–5 月 RONI=-0.1℃。"
            "产地降雨：加里曼丹中/西/东三省 6 月日均 <5mm，南苏门答腊 <2mm，多数主产区处历史同期偏低。"),
        "method": (
            "构建多指标气象监测框架：相对海表温度距平（Relative SSTA / RONI，相对海洋尼诺指数）、温跃层斜率、"
            "上层海洋热含量、850hPa 低层纬向风异常、MJO（Madden-Julian Oscillation，热带季内对流东传扰动）相位、"
            "OLR（Outgoing Longwave Radiation，向外长波辐射，反映对流强弱）与开尔文波追踪；给出西风爆发判据"
            "（≥+2m/s、单次≥10 天、3–6 月累计≥30 天）；产地按州做日均降雨跟踪（数据源 REFINITIV）。"),
        "borrow": (
            "提供了一套可直接映射到本项目特征工程的领先指标清单——把 RONI、温跃层斜率、上层热含量、西风异常、"
            "MJO 相位作为厄尔尼诺强度的先行信号，并明确“上层海洋热量领先海表温度”，正是 notebook 中 ONI 滞后特征"
            "的物理依据；分州日均降雨可细化 PRCP（降水）特征口径，并作为减产滞后验证的锚点。"),
    },
    {
        "no": "②",
        "title": "棕榈油：未来可能有较大幅度的上涨行情（专题报告）",
        "org": "五矿期货（农产品研究 · 杨泽元）",
        "date": "2026-06-22",
        "tag": "棕榈油专题 / 天气+价格",
        "link": NOTE_LINK,
        "conclusion": (
            "明确看多：假设国际原油不出现大幅下跌，在厄尔尼诺、印尼 B50 生柴政策、印尼政府出口管控三重因素推动下，"
            "棕榈油价格可能有较大幅度上涨；交易上推荐逢低买入 01、05 合约及 1–5 月差正套。并援引 NOAA 观点，"
            "本次厄尔尼诺可能成为自 1950 年有记录以来最严重的一次，是四份报告中对天气与价格最乐观、立场最鲜明的一份。"),
        "data": (
            "厄尔尼诺概率：5–7 月 97%、6–9 月 99%、8 月至明年 1 月 100%；11 月–明年 1 月超强厄尔尼诺概率升至 63%。"
            "截至 6/10 当周 Nino3.4 相对海温 +0.9℃（环比 +0.7，较 2015 同期 +0.6）。截至 6/15 累计降雨距平："
            "印尼 975mm(-33.96%)、马来 711mm(-32.25%)、印度 78mm(-53.54%)、泰国 184mm(-51.36%)。USDA：25/26 全球"
            "棕榈油产量 8127 万吨（印尼 4670+马来 2020，合占 82.31%），期末库存仅 1544 万吨（库消比 18.99%）；"
            "GAPKI 3 月底印尼库存 256.8 万吨、MPOB 5 月底马来 243 万吨；B50 较 B40 年增需求约 200 万吨。"),
        "method": (
            "自上而下的“供需+政策”驱动逻辑：以 NOAA 概率/强度预测定天气方向，用 Bloomberg 各国累计降雨距平印证"
            "干旱，引 USDA 产量/库存与 GAPKI、MPOB 月度库存量化供给偏紧，再叠加 B50 生柴政策（年增约 200 万吨消费）"
            "与印尼 DSI 独家出口管控（9/1 正式落地）推演需求增量与可流通货源收缩，最后落到具体合约与套利方向。"),
        "borrow": (
            "把“厄尔尼诺→东南亚降雨距平→棕榈油减产→低库存放大缺口→价格上行”的完整传导链做了量化拆解，恰好对应"
            "README 待补的 R1 传导链图；其多国降雨减幅百分比、库消比、政策需求增量可作为本项目价格预测模型的外生"
            "冲击变量与情景假设输入；并提供了可执行的合约与套利方向，便于将模型信号转化为交易校验。"),
    },
    {
        "no": "③",
        "title": "2026 年半年度策略报告（油脂）：分歧与共识，油脂势在何方？",
        "org": "中信期货（农业组）",
        "date": "2026-06-17",
        "tag": "油脂半年策略 / 厄尔尼诺与价格",
        "link": NOTE_LINK,
        "conclusion": (
            "上半年油脂在“现实宽松 vs 预期收紧”博弈下高位宽幅震荡；下半年围绕宏观货币、原油能源、印尼 B50 执行、"
            "厄尔尼诺天气四条主线运行。策略：单边逢低布局棕榈油、菜油多单，豆油区间偏弱对待；套利重点把握豆棕、"
            "菜豆价差区间机会。强弱格局判断为棕榈油 > 菜油 > 豆油 且长期固化，是四份中体系最完整、最强调多空分歧与风险的一份。"),
        "data": (
            "B50 全年刚性消耗约 1500 万吨棕榈油（占印尼产量 30–31%），较 B40 的 1300 万吨年增约 200 万吨（上半年 B40"
            "耗 650 万吨、下半年 B50 耗 850 万吨）。原油-政策阈值：布伦特 >75 美元/桶执行风险低、65–75 低风险、<65 中等风险。"
            "降雨—单产滞后 6–9 个月。策略价位：P2701 支撑 9400–9600、目标 9900–10200/上探 10200–10500 元/吨；"
            "豆油 8100–8300 支撑；菜油 9500–9600 支撑。价差区间：豆棕 -1400~-700、菜豆 -1800~-1100 元/吨。"),
        "method": (
            "采用“分歧—共识”框架梳理多空逻辑，配合原油—生物柴油—油脂价格传导机制；用 POGO/BOGO 价差"
            "（棕榈油、豆油与柴油价差，衡量生柴生产经济性）、美国 RIN-D4 碳积分价格判断生柴盈利；按棕榈树物候"
            "（花序分化/开花坐果/果实发育）解释减产滞后；用土壤墒情、干旱监测、NDVI（Normalized Difference"
            " Vegetation Index，归一化植被指数）评估全球各产区天气。"),
        "borrow": (
            "提供厄尔尼诺影响油脂的“结构性分化”全球视角（东南亚棕榈减产 vs 北美大豆影响有限 vs 加澳菜籽减产），"
            "尤其量化了“降雨—棕榈单产 6–9 个月滞后”机制——正是本项目 ONI_lag_{1,3,6,9,12} 滞后特征选择的产业依据；"
            "其价差区间与关键价位可作为模型输出结果的合理性校验区间，并提示需以“分歧/风险”制衡单边乐观判断。"),
    },
    {
        "no": "④",
        "title": "农业策略日报：天气顾虑需求乐观，棕榈油走势偏强",
        "org": "中信期货（农业组）",
        "date": "2026-04-22",
        "tag": "策略日报 / 厄尔尼诺主题起点",
        "link": NOTE_LINK,
        "conclusion": (
            "厄尔尼诺天气关注度提升叠加印尼下半年强制执行 B50 生柴掺混的文件，棕榈油领涨油脂、走势偏强。"
            "援引国家气候中心：预计 5 月进入厄尔尼诺状态、夏秋季形成中等及以上强度事件、至少持续至年底，"
            "但强调影响存在滞后、现在断言尚早而风险显著上升。展望豆油/棕油/菜油均“震荡偏强”，建议逢低买入，"
            "是四份中时间最早、定性为主的“主题起点”切片。"),
        "data": (
            "以定性判断与高频现货为主：国家气候中心给出 5 月进入厄尔尼诺、夏秋季中等及以上强度、持续至年底的"
            "时间与强度口径；天然橡胶等其他农产品同样以“强厄尔尼诺预期支撑盘面”定调；油脂板块统一展望为"
            "“震荡偏强”。列示风险因素：原油、地缘、天气异常、贸易关系、生物柴油政策、宏观等对油脂市场的波动风险。"),
        "method": (
            "采用多品种横向扫描的策略日报体例，对每个品种给出“信息/逻辑/展望/风险”四段式；判断方向上以"
            "宏观（原油、中东地缘）+ 政策（印尼 B50）+ 天气（厄尔尼诺）三因子驱动，配合现货报价、基差、"
            "持仓与仓单等高频数据做跟踪，强调风险因素清单以约束方向性结论。"),
        "borrow": (
            "时间上（4 月）早于其余三份，可作为厄尔尼诺交易主题的“起点切片”：与 6 月报告对比，能清晰观察预期"
            "如何被逐步计价（从“为时过早”→“风险显著上升”→“或为最强”），是预期演变的天然时间标尺；"
            "其“宏观+政策+天气”三因子框架可直接作为本项目模型特征分组（宏观面/政策事件/气象）的简洁模板。"),
    },
]


def build():
    doc = Document()
    # 默认正文字体
    normal = doc.styles["Normal"]
    normal.font.name = EAST_ASIA_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)

    # 标题
    add_para(doc, "券商研报精华萃取：棕榈油 × 厄尔尼诺",
             size=18, bold=True, font=HEAD_FONT,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
             color=(0x1F, 0x4E, 0x79))
    add_para(doc, "交付物 R3 / #4 · 单品种聚焦 Palm Oil（CPO=F） · 整理日期 2026-06-22",
             size=10, font=HEAD_FONT, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=2, color=(0x80, 0x80, 0x80))
    add_para(doc, "说明：以下要点均为对研报原文的转述与提炼，非原文复制粘贴；英文术语首次出现附中文释义。",
             size=9, font=EAST_ASIA_FONT, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=8, color=(0x80, 0x80, 0x80))

    # 一、清单表
    add_heading(doc, "一、研报清单表", size=13)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["#", "标题", "机构", "日期", "类型 / 链接"]
    widths = [0.4, 4.0, 1.7, 1.0, 1.6]
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        fill_cell(hdr[i], h, bold=True, size=9.5, color=(0xFF, 0xFF, 0xFF), font=HEAD_FONT)
        shade_cell(hdr[i], "1F4E79")
    for rp in REPORTS:
        cells = table.add_row().cells
        fill_cell(cells[0], rp["no"], bold=True, size=9.5)
        fill_cell(cells[1], rp["title"], size=9.5)
        fill_cell(cells[2], rp["org"], size=9.5)
        fill_cell(cells[3], rp["date"], size=9.5)
        fill_cell(cells[4], f'{rp["tag"]}\n（{rp["link"]}）', size=9)
    # 设置列宽
    from docx.shared import Inches
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    add_para(doc, "覆盖度自检：棕榈油专题 ≥1（①②均为专题）；天气/厄尔尼诺与农产品价格 ≥2（①②③④四份均涉及）。"
                  "注：交付标准建议清单 5 份，本次依据现有 4 份材料整理，可补充 1 份天气/价格类研报达 5 份口径。",
             size=9, color=(0x80, 0x80, 0x80), space_after=8)

    # 二、逐份萃取
    add_heading(doc, "二、逐份萃取（结论 / 数据指标 / 方法 / 可借鉴点）", size=13)
    for rp in REPORTS:
        add_heading(doc, f'{rp["no"]} {rp["org"]}《{rp["title"]}》（{rp["date"]}）',
                    size=11, color=(0x2E, 0x74, 0xB5), space_before=8)
        add_labeled(doc, "结论", rp["conclusion"])
        add_labeled(doc, "数据指标", rp["data"])
        add_labeled(doc, "方法", rp["method"])
        add_labeled(doc, "可借鉴点", rp["borrow"])

    # 三、交叉对比
    add_heading(doc, "三、四份研报交叉对比（≥300 字）", size=13)
    cross = [
        ("时间轴与预期演变",
         "四份报告沿“预期—实证—体系—宣告”递进：4 月中信日报（④）是主题起点，国家气候中心口径仍是"
         "“为时过早、风险显著上升”；6/10 东证周度（①）以 NOAA 多指标给出气象实证，确认热量储备超 2015 年；"
         "6/17 中信半年报（③）搭出完整策略体系并强调多空分歧；6/22 五矿专题（②）最为激进，直接喊出“或为有"
         "记录以来最强”。可见两个月内厄尔尼诺强度判断持续上修、价格预期被逐步计价。"),
        ("共识",
         "四份一致看多/偏强棕榈油，核心驱动高度一致——厄尔尼诺天气减产 + 印尼 B50 生柴需求增量 + 全球棕榈油"
         "低库存，且都判断棕榈油强于豆油。中信半年报与五矿专题进一步把驱动量化（B50 年增约 200 万吨、库消比"
         "约 19%），印证“生柴政策已取代食用刚需成为油脂第一定价因子”的判断。"),
        ("数据口径差异（引用须谨慎）",
         "同一指标因时点/口径不同而数值有别：厄尔尼诺 5–7 月概率，东证（5 月预测）为 82%，五矿（6 月预测）为 97%；"
         "Nino3.4 相对海温，东证 6/3 为 0.7℃、五矿 6/10 为 0.9℃。降雨口径亦不同：东证用产地分州日均 mm（REFINITIV），"
         "五矿用国别累计 mm 距平（Bloomberg）。本项目引用时必须标注日期、海区与数据源，避免混用。"),
        ("分歧与风险立场",
         "中信半年报（③）最为均衡，专设三大分歧——原油下行是否击穿 B50 财政兜底、厄尔尼诺减产实际幅度、"
         "加拿大菜籽利空是否已提前计价，并提示厄尔尼诺快速消退风险；五矿（②）立场最乐观，对回调与证伪风险着墨较少。"
         "二者并读可形成“乐观情景 + 风险对冲”的平衡视角。"),
        ("滞后性这一关键共识对建模的意义",
         "③强调“降雨—棕榈单产 6–9 个月滞后”、①强调“上层海洋热量领先海表温度”，共同指向：2026 年的天气冲击"
         "真正体现到产量要到 2027 年，且短期偏干反而利于鲜果串采收。这直接支持本项目以 ONI_lag_{1,3,6,9,12} 构造滞后"
         "特征，并提示模型对“近月偏弱、远月升水”的结构要有解释力。"),
        ("对本项目的统一启示",
         "①验证并细化 ONI/降水滞后特征；②把传导链量化为可输入的外生冲击变量；③的价差区间与关键价位作为模型"
         "输出的合理性校验带；④的“宏观+政策+天气”三因子可作为特征分组模板。B50 与印尼出口管控宜处理为政策"
         "哑变量（dummy variable，0/1 事件标记），多源降雨距平用于交叉校验，避免单一数据源偏差。"),
    ]
    for sub, body in cross:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        r1 = p.add_run(f"· {sub}：")
        set_cn_font(r1, font=HEAD_FONT, size=10.5, bold=True, color=(0x1F, 0x4E, 0x79))
        r2 = p.add_run(body)
        set_cn_font(r2, font=EAST_ASIA_FONT, size=10.5)

    doc.save(OUT_PATH)
    print("SAVED:", OUT_PATH)


if __name__ == "__main__":
    build()
