# -*- coding: utf-8 -*-
"""生成导师二次修改意见落实 Word 报告。

针对导师三条最新反馈：
  1. 样本窗口 132 个月太少，前推 1-2 年
  2. 气象因子之间同步关系不显著，给出原因和模型调整方案
  3. 8-10 月领先期显著可能是季节相位重合，检查 z-score 是否漏去掉季节性

输出:
  reports/Mentor_Revision_Summary_Report_v2.docx

说明:
  本脚本只读取既有研究总表和图表，不改变模型、不覆盖 Excel。
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
REPORTS_DIR = PROJECT_ROOT / "reports"
MODEL_DIR = PROJECT_ROOT / "code" / "model"
FIGURES_DIR = MODEL_DIR / "figures"
MASTER_XLSX = PROJECT_ROOT / "data" / "processed" / "archives" / "00_palm_oil_research_master_malaysia.xlsx"
OUT_DOCX = REPORTS_DIR / "Mentor_Revision_Summary_Report_v2.docx"


def add_para(doc: Document, text: str = "", *, bold: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10.5)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text).font.size = Pt(10.5)
    return p


def add_formula(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    return p


def make_table(doc: Document, headers: list[str], rows: list[list[object]], font_size: int = 9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(header))
        run.bold = True
        run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            cells[i].paragraphs[0].add_run(str(value)).font.size = Pt(font_size)
    return table


def add_picture_if_exists(doc: Document, path: Path, caption: str, width: float = 5.8):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = add_para(doc, caption, italic=True, size=9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"图表缺失: {path}", italic=True, size=9)


def fmt(x, digits: int = 4) -> str:
    try:
        if pd.isna(x):
            return "-"
        return f"{float(x):.{digits}f}"
    except (TypeError, ValueError):
        return str(x)


def read_inputs():
    master = pd.read_excel(MASTER_XLSX, sheet_name="02_Master_Monthly")
    season = pd.read_excel(MASTER_XLSX, sheet_name="03_Seasonality")
    zstats = pd.read_excel(MASTER_XLSX, sheet_name="04_ZScore_Stats")
    weather_corr = pd.read_excel(MASTER_XLSX, sheet_name="05_Weather_Corr")
    lag_corr = pd.read_excel(MASTER_XLSX, sheet_name="06_Lag_Corr")
    checks = pd.read_excel(MASTER_XLSX, sheet_name="01_Source_Check")
    return master, season, zstats, weather_corr, lag_corr, checks


def main() -> None:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    master, season, zstats, weather_corr, lag_corr, checks = read_inputs()

    n_backcast = int((master["Area_Source"] == "BACKCAST").sum()) if "Area_Source" in master.columns else 0
    n_mpob = int((master["Area_Source"] == "MPOB").sum()) if "Area_Source" in master.columns else len(master)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)

    title = doc.add_heading("导师二次修改意见落实报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = add_para(
        doc,
        "马来西亚棕榈油单产模型：样本扩展 + 去季节 z-score + 领前期季节相位诊断",
        italic=True, size=10,
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "数据口径：马来西亚棕榈油产量、成熟面积、ONI、马来西亚降水、马来西亚温度。", bold=True)
    add_para(doc, f"样本窗口：{master['Date'].min()} 至 {master['Date'].max()}，共 {len(master)} 个月（MPOB 实测 {n_mpob} 行 + 线性外推 {n_backcast} 行）。")

    # ─────────────────────────── 摘要 ───────────────────────────
    doc.add_heading("一、本次修订要点（结论先行）", level=1)
    add_bullet(doc, f"样本从 132 个月扩展到 {len(master)} 个月：把 MPOB 成熟面积 2015-2025 用线性外推补齐到 2010-2014，产量/气候数据原本就从 2007 起可用，因此样本起点提前到 2010-01。")
    add_bullet(doc, "z-score 升级为双口径：全序列 z-score（只去量纲，保留季节性）+ 按月份 z-score（去量纲 + 去季节性）。前者用于跨变量横截面比较，后者专门用于领先期分析。")
    add_bullet(doc, "领先期分析升级为三口径对比：原始值 / 全 z / 按月 z。结论非常关键——之前看到的 TAVG 领先 8 月 r=-0.70、PRCP 领先 10 月 r=+0.63 在去季节后几乎全部塌缩到 |r|<0.12，证实是季节相位重合造成的伪相关。")
    add_bullet(doc, "真正通过去季节检验的领先信号只有 PRCP_DEV_3m 领先 8-9 个月（去季节后 r≈0.15，p<0.05，斜率为正，符合农学逻辑）。")
    add_bullet(doc, "气象同步相关性整体偏弱，主因是马来西亚位于赤道季内振荡带，ONI 通过 Walker 环流影响本地降水/温度的传导路径在月度同步尺度上很弱，需要改用滞后或区域聚合方案。")

    # ─────────────────────────── 反馈 1 ───────────────────────────
    doc.add_heading("二、反馈 1：样本窗口扩展", level=1)
    add_para(doc, "导师反馈：132 个月样本偏少，前推 1-2 年，给样本外验证留足窗口。")
    add_para(doc, "数据现状排查：")
    add_bullet(doc, "产量（iFinD）：2007-01 ~ 2026-05，共 233 月，最早可用 2007。")
    add_bullet(doc, "马来西亚降水 / 温度（NASA POWER）：2007-01 ~ 2025-12，共 228 月。")
    add_bullet(doc, "ONI（NOAA）：1950-01 ~ 2026-04，共 916 月，远早于其他源。")
    add_bullet(doc, "成熟种植面积（MPOB）：仅 2015-2025，共 11 年——这才是原来卡在 132 行的瓶颈。")
    add_para(doc, "解决方案：MPOB 2015-2025 的 Mature_Area / Immature_Area / Total_Area 对 YEAR 做一元线性回归，外推 2010-2014 五年，并展开为月频。外推行在 Area_Source 列标记为 BACKCAST，便于后续敏感性分析时剔除或加权。")
    add_formula(doc, "Mature_Area(year) = a + b × year，a/b 由 2015-2025 最小二乘拟合得到")
    add_para(doc, "外推结果（年度值）：")
    backcast_yearly = (
        master[master["Area_Source"] == "BACKCAST"]
        .assign(YEAR=lambda d: pd.to_datetime(d["Date"]).dt.year)
        .groupby("YEAR", as_index=False)
        .agg(Mature_Area=("Mature_Area", "mean"), Total_Area=("Total_Area", "mean"))
    )
    make_table(
        doc,
        ["年份", "外推 Mature_Area (公顷)", "外推 Total_Area (公顷)"],
        [[int(r["YEAR"]), f"{r['Mature_Area']:.0f}", f"{r['Total_Area']:.0f}"] for _, r in backcast_yearly.iterrows()],
    )
    add_para(doc, "样本窗口变化：原 2015-01 ~ 2025-12 (132 月) → 现 2010-01 ~ 2025-12 (192 月)。新增 60 月，样本外窗口（最后 24 月）相对剩余 168 月训练样本，验证更稳健。")
    add_para(doc, "交叉核验：data_pipeline.build_dataset() 与主宽表的关键统计量逐项核对，33/33 项通过（差值 < 1e-4），证明外推与合并逻辑无误。")
    add_para(doc, "局限与后续改进：线性外推假设面积趋势延续，未考虑 2010-2014 马来西亚实际种植扩张节奏可能非线性。若导师能从 MPOB 历史 PDF 中补回 2010-2014 真实值，可替换 BACKCAST 行消除该假设。")

    # ─────────────────────────── 反馈 2 ───────────────────────────
    doc.add_heading("三、反馈 2：气象同步关系不显著的原因与模型调整方案", level=1)
    add_para(doc, "导师反馈：天气因子之间同步关系不显著，给出可能的原因和模型调整方案。")
    add_para(doc, "实测结果（双口径对比）：")
    make_table(
        doc,
        ["X", "Y", "原始 r", "原始显著性", "去季节 r", "去季节显著性"],
        weather_corr[["X变量", "Y变量", "原始_Pearson_r", "原始_显著性",
                      "去季节_Pearson_r", "去季节_显著性"]].values.tolist(),
        font_size=8,
    )
    add_para(doc, "观察到的事实：")
    add_bullet(doc, "ONI_lag12 vs PRCP_DEV_3m 同步 r=-0.04 不显著；ONI_lag12 vs TAVG_DEV_3m 去季节后 r=+0.23 显著；其余均不显著。")
    add_bullet(doc, "PRCP vs TAVG 原始 r=-0.19 (p<0.05) 显著，但去季节后 r=-0.12 不显著——这再次说明原始显著主要由季节相位驱动。")
    add_para(doc, "可能原因（按重要性排序）：")
    add_number(doc, "尺度不匹配：ONI 反映赤道太平洋海温异常，是大尺度遥相关信号；马来西亚本地月度降水/温度受季内振荡 (MJO)、ITCZ 季节迁移、地形等多重局域因素调制，月度同步尺度上信号被淹没。")
    add_number(doc, "时间延迟：ENSO 通过大气环流传导到东南亚降水通常有 1-3 月滞后，同步 (lag=0) 自然弱；这与后面领先期分析中 PRCP_DEV_3m 在 8-9 月前对单产显著相符——ONI 的真实影响通过滞后链路释放。")
    add_number(doc, "样本量不足：去季节后有效样本 ~190 月，月度同步散点噪声大；如能拿到日/周频数据并聚合到季度，相关结构会更清晰。")
    add_number(doc, "区域聚合不足：用全国平均单点降水/温度掩盖了东马（沙巴/砂拉越）与西马气候差异，相关信号被平均稀释。")
    add_para(doc, "模型调整方案：")
    add_bullet(doc, '方案 A（推荐）：保留 ONI_lag12 + PRCP_DEV_3m + TAVG_DEV_3m 的当前结构，因为它们已经天然形成"ONI 背景信号 → 滞后降水距平 → 滞后温度距平 → 单产"的传导链，弱同步相关反而说明三者近似正交，共线性低，适合多元回归。')
    add_bullet(doc, "方案 B：把 ONI 改为 ONI_lag3 或 ONI_lag6 以匹配传导延迟；在 train_model.py 中加入 ONI 滞后网格搜索。")
    add_bullet(doc, "方案 C：引入区域分层数据（西马 / 沙巴 / 砂拉越），用面板回归 (panel regression) 替代全国聚合，捕捉空间异质性。")
    add_bullet(doc, "方案 D：把月度数据聚合到季度，降低单月噪声，重新评估同步相关。")

    # ─────────────────────────── 反馈 3 ───────────────────────────
    doc.add_heading("四、反馈 3：季节相位重合诊断与 z-score 修复", level=1)
    add_para(doc, "导师反馈：8-10 月领先期显著可能是季节相位重合；检查第 3 步 z-score 是否漏掉关键步骤、没把季节性去掉。")
    add_para(doc, "诊断结论：导师判断完全正确。原版 z-score 用全序列均值/标准差做标准化，只去量纲、保留季节性。棕榈油单产和温度都有强年内周期，二者季节相位若错开 8-10 个月，就会在原始序列上呈现强“领先相关”——但这只是日历周期重合，不是真因果。")

    doc.add_heading("4.1 z-score 修复：双口径并列", level=2)
    add_formula(doc, "全序列 z (原版，只去量纲):  z = (x - mean(x全序列)) / std(x全序列)")
    add_formula(doc, "按月份 z (新版，去量纲+去季节):  z_m = (x - mean(x同月历史)) / std(x同月历史)")
    add_para(doc, "按月份 z-score 等价于“在同一个月内做标准化”——比如所有 8 月的产量和温度各自做标准化，季节性均值被扣到 0，剩下的纯是相对该月常态的偏离。这一步是原版漏掉的关键。")
    add_para(doc, "z-score 统计量对比：")
    make_table(
        doc,
        ["变量", "全序列均值", "全序列标准差", "全z均值", "全z标准差", "按月z均值", "按月z标准差"],
        zstats[["变量", "全序列均值", "全序列标准差", "全z均值", "全z标准差",
                "按月z均值", "按月z标准差"]].values.tolist(),
        font_size=8,
    )

    doc.add_heading("4.2 领先期三口径对比：核心证据", level=2)
    add_para(doc, "对每个 (变量, k) 同时给出原始 r、全 z r、按月 z r，并标注显著性。判定规则：")
    add_bullet(doc, "真信号 = 按月 z 列仍 |r| 较大且 p<0.05 的领先期。")
    add_bullet(doc, "伪相关 = 原始显著但按月 z 列塌缩到 |r|<0.12 且不显著的领先期。")

    add_para(doc, "伪相关典型案例（原版误判为“显著领先信号”）：")
    fake_rows = []
    for _, r in lag_corr.iterrows():
        if r["原始_显著性"] in ("***", "**", "*") and r["按月z_显著性"] not in ("***", "**", "*"):
            fake_rows.append([
                r["天气变量"], int(r["领先月数k"]),
                fmt(r["原始_r"], 4), r["原始_显著性"],
                fmt(r["按月z_r(去季节)"], 4), r["按月z_显著性"],
            ])
    make_table(doc, ["变量", "k", "原始 r", "原始显著性", "去季节 r", "去季节显著性"], fake_rows[:15], font_size=8)

    add_para(doc, "真信号案例（去季节后仍显著）：")
    real_rows = []
    for _, r in lag_corr.iterrows():
        if r["按月z_显著性"] in ("***", "**", "*"):
            real_rows.append([
                r["天气变量"], int(r["领先月数k"]),
                fmt(r["原始_r"], 4), r["原始_显著性"],
                fmt(r["按月z_r(去季节)"], 4), r["按月z_显著性"],
                fmt(r["去季节表示函数_斜率b"], 4),
            ])
    make_table(doc, ["变量", "k", "原始 r", "原始显著性", "去季节 r", "去季节显著性", "去季节斜率 b"], real_rows, font_size=8)

    add_para(doc, "关键解读：")
    add_bullet(doc, "TAVG vs Yield 领先 8 月：原始 r=-0.7014 极强显著，去季节后 r=-0.0483 几乎归零——这是季节相位重合的典型伪相关。")
    add_bullet(doc, "PRCP vs Yield 领先 10 月：原始 r=+0.6281 强显著，去季节后 r=+0.1096 不显著——同样是季节相位伪相关。")
    add_bullet(doc, "PRCP_DEV_3m vs Yield 领先 8-9 月：去季节后仍 r≈0.15、p<0.05、斜率 +0.21~+0.23（正向，符合农学逻辑：3 月滚动降水偏多→8-9 月后单产偏高）——这是唯一通过季节相位检验的真信号。")
    add_bullet(doc, "ONI_lag12 在去季节后 r 反而从 0.21 升到 0.30，说明 ONI 信号是真气候信号而非季节相位。")

    add_picture_if_exists(doc, FIGURES_DIR / "lag_correlation_TAVG.png",
                          "图 1：温度领先 1-12 期三口径对比。蓝色=原始 r，橙色=全 z r，红色=按月 z r（去季节）。可见蓝色在 k=8 处冲到 -0.70，红色在全部 k 上几乎为 0——季节相位伪相关。")
    add_picture_if_exists(doc, FIGURES_DIR / "lag_correlation_PRCP.png",
                          "图 2：降水领先 1-12 期三口径对比。k=10 处蓝色 +0.63，红色仅 +0.11，同样为伪相关。")
    add_picture_if_exists(doc, FIGURES_DIR / "lag_correlation_PRCP_DEV_3m.png",
                          "图 3：3 月滚动降水距平领先 1-12 期三口径对比。k=8-9 处红色仍稳定在 +0.15 附近，且 p<0.05——真信号。")
    add_picture_if_exists(doc, FIGURES_DIR / "lag_correlation_ONI_lag12.png",
                          "图 4：ONI_lag12 领先 1-12 期三口径对比。红色去季节后反而增强，证明 ONI 是真气候信号。")

    # ─────────────────────────── 方法 ───────────────────────────
    doc.add_heading("五、计算方法与公式汇总", level=1)
    add_number(doc, "数据合并：产量、气候、面积按 YYYY-MM 月度对齐 join，面积由年频展开为月频（外推部分标记 BACKCAST）。")
    add_formula(doc, "Yield_t = Production_t / Mature_Area_t")
    add_number(doc, "气候距平：减去同月历史平均，得到对季节常态的偏离。")
    add_formula(doc, "PRCP_DEV_t = PRCP_t - mean(PRCP_same_month)")
    add_formula(doc, "TAVG_DEV_t = TAVG_t - mean(TAVG_same_month)")
    add_number(doc, "3 月滚动距平：降低单月噪声。")
    add_formula(doc, "PRCP_DEV_3m_t = mean(PRCP_DEV_t..t-2)")
    add_number(doc, "z-score 双口径：")
    add_formula(doc, "全序列 z:  z = (x - mean(x全序列)) / std(x全序列)  # 只去量纲")
    add_formula(doc, "按月份 z:  z_m = (x - mean(x同月)) / std(x同月)      # 去量纲 + 去季节")
    add_number(doc, "领先期三口径：weather(t-k) 对 Yield(t)，k=1..12，分别用原始/全 z/按月 z 求 Pearson r 与 p。")
    add_number(doc, "表示函数：一元线性回归 Yield_mz = a + b × weather_mz(t-k)，仅对去季节口径计算，记录 r/p/样本数。")

    # ─────────────────────────── 与预测模型关系 ───────────────────────────
    doc.add_heading("六、与现有预测模型的关系及下一步建议", level=1)
    add_para(doc, "本次修订是研究验证层，不直接覆盖网站 predict.py 的运行。现有预测模型结构：")
    add_formula(doc, "Yield_pred = α + β_trend×Trend + β_ONI×ONI_lag12 + β_PRCP×PRCP_DEV_3m + β_TAVG×TAVG_DEV_3m + β_month")
    add_formula(doc, "Production_pred = Yield_pred × Latest_Mature_Area")
    add_para(doc, "本次发现对现有模型的启示：")
    add_bullet(doc, "现有模型使用 PRCP_DEV_3m / TAVG_DEV_3m / ONI_lag12 是合理的——这三个变量已天然去季节（距平本身就是减去同月均值），且领先期分析显示 PRCP_DEV_3m 在去季节后仍显著。")
    add_bullet(doc, "现有模型使用同步 (lag=0) 的 PRCP_DEV_3m / TAVG_DEV_3m，但研究层发现 PRCP_DEV_3m 真实领先期是 8-9 月。下一步可在 train_model.py 中尝试 PRCP_DEV_3m_lag8/9，看样本外 RMSE 是否改善。")
    add_bullet(doc, "原版模型中的月份季节项 (β_month) 是保留季节性的正确做法——保留季节性做建模，但在做相关分析时用按月份 z-score 去季节，二者目的不同。")
    add_bullet(doc, "面积外推 (BACKCAST 行) 会进入训练样本，建议下一步在 train_model.py 增加“仅 MPOB”和“含 BACKCAST”两个训练版本对比，验证外推假设对模型稳定性的影响。")

    doc.add_heading("七、最终结论", level=1)
    add_para(doc, "导师三条反馈全部落实并有统计证据支撑：")
    add_number(doc, "样本窗口从 132 月扩展到 192 月（前推 5 年到 2010-01），通过面积线性外推实现；交叉核验 33/33 通过。")
    add_number(doc, "气象同步关系弱的原因被定位为尺度不匹配 + 传导延迟 + 区域聚合稀释；推荐方案 A（保留当前近似正交的三因子结构）。")
    add_number(doc, "原版 z-score 确实只去量纲未去季节，已升级为双口径；领先期三口径对比证实 TAVG/PRCP 的 8-10 月领先期是季节相位伪相关，唯一真信号是 PRCP_DEV_3m 领先 8-9 月。")

    doc.save(OUT_DOCX)
    print(f"[OK] Saved Word report: {OUT_DOCX}")


if __name__ == "__main__":
    main()
