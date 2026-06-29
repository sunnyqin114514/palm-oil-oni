# -*- coding: utf-8 -*-
"""生成导师修改意见落实 Word 报告。

输出:
  reports/Mentor_Revision_Summary_Report.docx

说明:
  本脚本只读取既有研究总表和图表，不改变模型、不覆盖 Excel。
"""
from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
REPORTS_DIR = PROJECT_ROOT / "reports"
MODEL_DIR = PROJECT_ROOT / "code" / "model"
FIGURES_DIR = MODEL_DIR / "figures"
MASTER_XLSX = PROJECT_ROOT / "data" / "processed" / "archives" / "00_palm_oil_research_master_malaysia.xlsx"
OUT_DOCX = REPORTS_DIR / "Mentor_Revision_Summary_Report.docx"


def add_para(doc: Document, text: str = "", *, bold: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10.5)
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text).font.size = Pt(10.5)
    return p


def add_formula(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    return p


def make_table(doc: Document, headers: list[str], rows: list[list[object]], font_size: int = 9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(header))
        run.bold = True
        run.font.size = Pt(font_size)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            cells[i].paragraphs[0].add_run(str(value)).font.size = Pt(font_size)
    return table


def add_picture_if_exists(doc: Document, path: Path, caption: str, width: float = 5.8):
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = add_para(doc, caption, italic=True, size=9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"图表缺失: {path}", italic=True, size=9)


def fmt(x: float, digits: int = 4) -> str:
    return f"{float(x):.{digits}f}"


def read_inputs():
    master = pd.read_excel(MASTER_XLSX, sheet_name="02_Master_Monthly")
    season = pd.read_excel(MASTER_XLSX, sheet_name="03_Seasonality")
    zstats = pd.read_excel(MASTER_XLSX, sheet_name="04_ZScore_Stats")
    weather_corr = pd.read_excel(MASTER_XLSX, sheet_name="05_Weather_Corr")
    lag_corr = pd.read_excel(MASTER_XLSX, sheet_name="06_Lag_Corr")
    checks = pd.read_excel(MASTER_XLSX, sheet_name="01_Source_Check")
    return master, season, zstats, weather_corr, lag_corr, checks


def main() -> None:
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    master, season, zstats, weather_corr, lag_corr, checks = read_inputs()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)

    title = doc.add_heading("导师修改意见落实报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = add_para(doc, "马来西亚棕榈油单产模型：数据总表、季节性、标准化、相关性与领先期分析", italic=True, size=10)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(doc, "数据口径：马来西亚棕榈油产量、成熟面积、ONI、马来西亚降水、马来西亚温度。", bold=True)
    add_para(doc, f"样本窗口：{master['Date'].min()} 至 {master['Date'].max()}，共 {len(master)} 个月。")

    doc.add_heading("一、结论先行", level=1)
    add_bullet(doc, f"研究总表已生成，并对数据源做交叉核验：{int((checks['通过'] == '✓').sum())}/{len(checks)} 项通过。")
    add_bullet(doc, "单产有明显季节性：9-10 月偏高，1-2 月偏低，因此后续预测必须保留月份季节项。")
    add_bullet(doc, "z-score 标准化已完成：不同单位的变量（吨/公顷、毫米、摄氏度、ONI）被转换到同一尺度，便于比较。")
    add_bullet(doc, "气象变量之间同步相关性较弱，说明 ONI、降水、温度各自提供不同信息。")
    add_bullet(doc, "领先期分析显示：降水在 9-12 月前对单产有正向信号，温度在 7-9 月前有明显负向信号。")

    doc.add_heading("二、数据逻辑与因果推理链", level=1)
    add_para(doc, "本项目不是直接用天气预测产量吨数，而是先剥离出单产，再把天气解释为单产变化的一部分。这样可以避免成熟面积扩张把天气信号掩盖掉。")
    add_number(doc, "先用产量和成熟面积计算单产。")
    add_formula(doc, "Yield_t = Production_t / Mature_Area_t")
    add_number(doc, "再把降水和温度转换成“相对同月正常水平的偏离”。")
    add_formula(doc, "PRCP_DEV_t = PRCP_t - mean(PRCP_same_month)")
    add_formula(doc, "TAVG_DEV_t = TAVG_t - mean(TAVG_same_month)")
    add_number(doc, "为了降低单月噪声，再计算 3 个月滚动距平。")
    add_formula(doc, "PRCP_DEV_3m_t = mean(PRCP_DEV_t, PRCP_DEV_{t-1}, PRCP_DEV_{t-2})")
    add_formula(doc, "TAVG_DEV_3m_t = mean(TAVG_DEV_t, TAVG_DEV_{t-1}, TAVG_DEV_{t-2})")
    add_number(doc, "最后用 z-score 消除变量单位差异。")
    add_formula(doc, "z_x = (x - mean(x)) / std(x)")
    add_para(doc, "因此，逻辑链为：历史天气与 ONI → 气候距平与标准化 → 单产波动 → 总产量还原。")

    doc.add_heading("三、具体数学运算示例", level=1)
    row = master.loc[master["Date"] == "2025-10"].iloc[0] if (master["Date"] == "2025-10").any() else master.iloc[-1]
    add_para(doc, f"以下用 {row['Date']} 的实际历史样本说明计算过程。")
    add_formula(doc, f"Yield = Production / Mature_Area = {row['Production']:.0f} / {row['Mature_Area']:.0f} = {row['Yield']:.6f} 吨/公顷")
    add_formula(doc, f"PRCP_DEV = PRCP - PRCP_CLIM = {row['PRCP']:.4f} - {row['PRCP_CLIM']:.4f} = {row['PRCP_DEV']:.4f} mm")
    add_formula(doc, f"TAVG_DEV = TAVG - TAVG_CLIM = {row['TAVG']:.4f} - {row['TAVG_CLIM']:.4f} = {row['TAVG_DEV']:.4f} °C")
    add_formula(doc, f"PRCP_DEV_3m = {row['PRCP_DEV_3m']:.4f}；TAVG_DEV_3m = {row['TAVG_DEV_3m']:.4f}")

    z_yield = zstats[zstats["变量"] == "Yield"].iloc[0]
    add_para(doc, "z-score 例子：")
    add_formula(doc, f"Yield_z = (Yield - mean(Yield)) / std(Yield) = ({row['Yield']:.6f} - {z_yield['原始均值']:.6f}) / {z_yield['原始标准差']:.6f}")

    doc.add_heading("四、研究总表 Excel 概括", level=1)
    add_para(doc, "Excel 总表只作为数据底稿，不在报告中铺开大表。它包含 8 个 Sheet：源文件核验、月度主宽表、季节性统计、z-score 标准化、气象相关性、领先期相关性和方法说明。")
    add_para(doc, f"文件路径：{MASTER_XLSX}")
    add_para(doc, "关键意义：导师可以直接打开 Excel 复核每个变量来源、计算列和统计结论；模型脚本也可复现该总表。")

    doc.add_heading("五、季节性特征", level=1)
    add_para(doc, "季节性不是凭图感判断，而是按月份 1-12 对每个变量计算均值、中位数和标准差。若某变量每年固定月份高、固定月份低，就说明它有季节性。")
    yield_s = season[season["变量"] == "Yield"].sort_values("均值", ascending=False)
    make_table(
        doc,
        ["变量", "高峰月份", "高峰均值", "低谷月份", "低谷均值", "解释"],
        [
            ["Yield 单产", int(yield_s.iloc[0]["月份"]), fmt(yield_s.iloc[0]["均值"], 4),
             int(yield_s.iloc[-1]["月份"]), fmt(yield_s.iloc[-1]["均值"], 4), "单产季节性强，模型必须保留月份项"],
            ["PRCP 降水", "10-11", "雨季偏高", "1-2", "相对偏低", "水分供给具有年内周期"],
            ["TAVG 温度", "4-5", "略高", "12-1", "略低", "温度季节性弱于产量和降水"],
        ],
    )
    add_picture_if_exists(doc, FIGURES_DIR / "seasonality_Yield.png", "图 1：单产 Yield 的季节性剖面。")
    add_picture_if_exists(doc, FIGURES_DIR / "seasonality_PRCP.png", "图 2：马来西亚降水 PRCP 的季节性剖面。")
    add_picture_if_exists(doc, FIGURES_DIR / "seasonality_TAVG.png", "图 3：马来西亚温度 TAVG 的季节性剖面。")

    doc.add_heading("六、气象同步相关性与表示函数", level=1)
    add_para(doc, "同步相关性回答的问题是：同一月份里，ONI、降水、温度之间是否互相强烈绑定。如果它们高度相关，多元模型会有共线性风险；如果相关性弱，则它们可分别提供信息。")
    add_formula(doc, "Pearson r = cov(X, Y) / (std(X) × std(Y))")
    add_formula(doc, "一元表示函数：Y = a + bX")
    make_table(
        doc,
        ["X", "Y", "r", "p值", "表示函数"],
        weather_corr[["X变量", "Y变量", "Pearson_r", "p值", "表达式"]].values.tolist(),
        font_size=8,
    )
    add_para(doc, "结论：所有同步相关性均不强，说明 ONI、降水、温度之间不是简单重复变量，保留在多元模型中有意义。")

    doc.add_heading("七、天气到单产的领先期关系", level=1)
    add_para(doc, "领先期分析回答的问题是：天气先发生，单产后变化，中间隔几个月最明显。计算方式是把天气变量向后平移 k 个月，用 weather(t-k) 对应 Yield(t)，k 从 1 到 12。")
    add_formula(doc, "Yield_t = a_k + b_k × Weather_{t-k}")
    add_formula(doc, "对每个 k 分别计算 Pearson r、p 值、斜率 b_k 和截距 a_k。")

    sig = lag_corr[lag_corr["p值"] < 0.05].copy()
    top_rows = []
    for var in ["PRCP", "TAVG", "ONI_lag12", "PRCP_DEV_3m", "TAVG_DEV_3m"]:
        sub = sig[sig["天气变量"] == var].copy()
        if sub.empty:
            continue
        sub["abs_r"] = sub["Pearson_r"].abs()
        best = sub.sort_values("abs_r", ascending=False).iloc[0]
        top_rows.append([best["天气变量"], int(best["领先月数k"]), best["Pearson_r"], best["p值"], best["表达式"]])
    make_table(doc, ["天气变量", "最佳领先月", "r", "p值", "表示函数"], top_rows, font_size=8)
    add_para(doc, "重点结论：")
    add_bullet(doc, "降水 PRCP 在 t-10 月与单产关系最强，r=+0.60，说明长期水分条件对后续果穗发育可能更关键。")
    add_bullet(doc, "温度 TAVG 在 t-8 月与单产关系最强，r=-0.69，说明较早期热胁迫可能压制后续单产。")
    add_bullet(doc, "ONI 的直接相关较弱，应更多理解为通过降水和温度传导的背景气候信号。")
    add_picture_if_exists(doc, FIGURES_DIR / "lag_correlation_PRCP.png", "图 4：降水领先 1-12 期与单产的相关性。")
    add_picture_if_exists(doc, FIGURES_DIR / "lag_correlation_TAVG.png", "图 5：温度领先 1-12 期与单产的相关性。")
    add_picture_if_exists(doc, FIGURES_DIR / "lag_correlation_ONI_lag12.png", "图 6：ONI 领先关系与单产的相关性。")

    doc.add_heading("八、与现有预测模型的关系", level=1)
    add_para(doc, "这次修订是研究验证层，不直接覆盖网站预测模型。现有预测模型已包含：趋势项、月份季节项、ONI_lag12、3 月滚动降水距平和 3 月滚动温度距平。")
    add_para(doc, "当前网站预测公式为：")
    add_formula(doc, "Yield_pred = α + β_trend×Trend + β_ONI×ONI_lag12 + β_PRCP×PRCP_DEV_3m + β_TAVG×TAVG_DEV_3m + β_month")
    add_formula(doc, "Production_pred = Yield_pred × Latest_Mature_Area")
    add_para(doc, "如果导师认可 PRCP(t-10) 和 TAVG(t-8) 的强领先信号，下一步才应进入模型重训：扩展 validate_leads.py 的搜索范围，并同步更新 predict.py 的滞后读取逻辑。")

    doc.add_heading("九、最终结论", level=1)
    add_para(doc, "本次修改把导师意见转化为可核验的数据底稿和统计证据：数据已统一，季节性已量化，量纲已消除，气象变量内部关系已检查，天气到单产的领先关系已系统扫描。")
    add_para(doc, "在研究结论层面，单产季节性最强；降水和温度的领先效应明显；ONI 更适合作为背景气候指数而非单独解释变量。")

    doc.save(OUT_DOCX)
    print(f"[OK] Saved Word report: {OUT_DOCX}")


if __name__ == "__main__":
    main()
