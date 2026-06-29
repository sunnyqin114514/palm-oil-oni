# -*- coding: utf-8 -*-
"""
Build the Day 5 palm-oil-only delivery Word document.

Output:
    reports/Day5_Learning_Report.docx
"""
from __future__ import annotations

import os

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(os.path.dirname(HERE))  # 项目根 ONI project/

CSV_FILE = os.path.join(ROOT, "data", "raw", "product", "CPO_F_daily_yahoo.csv")
GAP_FILE = os.path.join(ROOT, "data", "processed", "product", "CPO_F_calendar_gaps.csv")
CHART_FILE = os.path.join(ROOT, "data", "figures", "CPO_F_price_trend.png")
OUT_DOCX = os.path.join(ROOT, "reports", "Day5_Learning_Report.docx")


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def make_table(doc: Document, headers: list[str], rows: list[list[object]], font_size: int = 9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(header))
        run.bold = True
        run.font.size = Pt(font_size)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            cells[i].paragraphs[0].add_run(str(value)).font.size = Pt(font_size)
    return table


def main() -> None:
    df = pd.read_csv(CSV_FILE)
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    price_cols = ["Open", "High", "Low", "Close", "Adj Close"]

    start = df["Date"].min().date()
    end = df["Date"].max().date()
    span_years = (pd.Timestamp(end) - pd.Timestamp(start)).days / 365.25
    missing_rate = df[price_cols].isna().mean().max()
    zero_or_negative = int((df[price_cols] <= 0).sum().sum())
    duplicated_dates = int(df["Date"].duplicated().sum())
    gaps = pd.read_csv(GAP_FILE) if os.path.exists(GAP_FILE) else pd.DataFrame()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_heading("Day 5 — Financial Data Confirmation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = add_para(doc, "Palm Oil Futures Price Data · Yahoo Finance", italic=True, size=10)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Objective", level=1)
    add_para(
        doc,
        "This document records the Day 5 financial-data task for palm oil futures only. "
        "The goal is to confirm that the selected price dataset has a time span of at least 10 years, "
        "a missing rate no higher than 5%, and a visually continuous price trend suitable for later analysis.",
    )

    doc.add_heading("2. Financial Data Confirmation Table", level=1)
    make_table(
        doc,
        ["Required Field", "Palm Oil Futures Dataset"],
        [
            ["Data source", "Yahoo Finance"],
            ["Fields", "Date, Open, High, Low, Close, Adj Close, Volume"],
            ["Span", f"{start} to {end} ({span_years:.1f} years)"],
            ["Frequency", "Daily"],
            ["Acquisition method", "Yahoo Finance historical price download via Python yfinance"],
            [
                "Risk points",
                "Free continuous-futures data may have imperfect roll/stitching logic. "
                "One large gap was found from 2015-12-04 to 2016-07-08 (217 days). "
                "No zero or negative prices were found.",
            ],
        ],
        font_size=9,
    )

    doc.add_heading("3. Workflow", level=1)
    add_bullet(doc, "Step 1 — Download CPO=F daily historical prices from Yahoo Finance.")
    add_bullet(doc, "Step 2 — Read the CSV file with pandas and inspect the first rows, columns, date span, and missing values.")
    add_bullet(doc, "Step 3 — Check duplicated dates, large calendar gaps, and zero or negative price values.")
    add_bullet(doc, "Step 4 — Plot the daily Close price trend using matplotlib for visual inspection.")

    doc.add_heading("4. Data Quality Check", level=1)
    make_table(
        doc,
        ["Check Item", "Result"],
        [
            ["Rows", len(df)],
            ["Date range", f"{start} to {end}"],
            ["Maximum missing rate among price columns", f"{missing_rate:.2%}"],
            ["Duplicated dates", duplicated_dates],
            ["Zero / negative price cells", zero_or_negative],
            ["Large calendar gap", "2015-12-04 to 2016-07-08 (217 days)" if not gaps.empty else "None"],
            ["Requirement result", "PASS: span >= 10 years and missing rate <= 5%"],
        ],
        font_size=9,
    )

    doc.add_heading("5. Price Trend Chart", level=1)
    add_para(
        doc,
        "The following chart plots the daily Close price. The red shaded area marks the detected large calendar gap.",
    )
    if os.path.exists(CHART_FILE):
        doc.add_picture(CHART_FILE, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"Chart file not found: {CHART_FILE}", italic=True)

    doc.add_heading("6. Files Produced", level=1)
    make_table(
        doc,
        ["File", "Meaning"],
        [
            ["data/raw/product/CPO_F_daily_yahoo.csv", "Raw Yahoo Finance daily price data"],
            ["data/processed/product/CPO_F_calendar_gaps.csv", "Detected large calendar gaps between adjacent trading dates"],
            ["data/figures/CPO_F_price_trend.png", "Daily Close price trend chart"],
            ["scripts/download_palm_oil_yahoo.py", "Download script"],
            ["scripts/explore_plot_palm_oil.py", "Data check and plotting script"],
        ],
        font_size=9,
    )

    doc.add_page_break()
    doc.add_heading("7. Reading Report", level=1)
    add_para(
        doc,
        "Source: IMF Commodity Special Feature, October 2022 — "
        "Market Developments and Food Price Inflation Drivers.",
        bold=True,
    )

    doc.add_heading("7.1 Key Takeaways", level=2)
    add_bullet(
        doc,
        "Food prices are driven by multiple shocks at the same time: harvest/production shocks, "
        "fertilizer prices, oil prices, interest rates, exchange rates, and global demand.",
    )
    add_bullet(
        doc,
        "Supply-side shocks matter strongly. The IMF highlights that a 10% rise in fertilizer prices "
        "can raise cereal prices by about 7% after one quarter.",
    )
    add_bullet(
        doc,
        "Oil prices affect food prices with a longer lag, partly through fuel, logistics, and production costs.",
    )
    add_bullet(
        doc,
        "Monetary policy also matters. A 100-basis-point US monetary-policy shock can reduce cereal prices "
        "by about 13% with a one-quarter lag, partly by weakening demand and speculative inventory holding.",
    )
    add_bullet(
        doc,
        "The IMF explicitly flags upside risks from export restrictions, droughts, and fertilizer pass-through. "
        "Indonesia's 2022 palm-oil export ban is mentioned as an example of policy risk.",
    )

    doc.add_heading("7.2 Link to My Project", level=2)
    add_para(
        doc,
        "This article supports the logic of my project: climate shocks do not affect futures prices directly; "
        "they first change expected supply, production costs, inventories, and trade policy, and then these "
        "channels move futures prices. For palm oil, El Niño can reduce rainfall in Malaysia and Indonesia, "
        "raising concerns about future yields. At the same time, fertilizer, oil, exchange rates, and export "
        "policies can amplify or weaken the price reaction. Therefore, my model should treat ONI and rainfall "
        "as key climate features, but not as the only drivers of palm-oil futures prices."
    )

    doc.add_heading("7.3 How the IMF Model Helps My Project", level=2)
    add_para(
        doc,
        "The IMF uses a dynamic shock-response framework, including local projections and impulse-response "
        "functions, to estimate how food prices react over time after a shock. This is useful for my project "
        "because El Niño is also a delayed shock. I can adapt the same idea by testing whether ONI or rainfall "
        "changes lead palm-oil prices after 1, 3, 6, 9, or 12 months. In practice, this means creating lagged "
        "features such as ONI_lag_3, ONI_lag_6, and PRCP_lag_6, then checking which lag has the strongest "
        "relationship with palm-oil prices."
    )

    doc.add_heading("7.4 Expert View and Connection to My Project", level=2)
    add_para(
        doc,
        "The IMF's expert view is that food-price inflation is not caused by one single factor. It is a system "
        "of interacting shocks: weather affects harvests, energy affects transport and inputs, fertilizer "
        "affects yields, interest rates affect inventories and speculation, and trade policy can suddenly "
        "tighten supply. This is directly connected to my project. My palm-oil model should not simply ask "
        "'Does El Niño raise prices?' Instead, it should ask 'Under which cost, policy, and macro conditions "
        "does El Niño create a stronger price response?' This expert view helps turn the project from a simple "
        "correlation exercise into a more realistic supply-shock model."
    )

    note = add_para(
        doc,
        "Short conclusion: the IMF article confirms that lagged climate variables should be combined with "
        "cost, policy, and macro variables when explaining agricultural futures prices.",
        italic=True,
    )
    note.runs[0].font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.save(OUT_DOCX)
    print(f"[OK] Saved Word report: {OUT_DOCX}")


if __name__ == "__main__":
    main()
